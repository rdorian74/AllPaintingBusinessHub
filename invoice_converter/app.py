#!/usr/bin/env python3
"""
Invoice to QuickBooks CSV Converter
Uses Claude AI for intelligent invoice data extraction
"""

import os
import re
import csv
import json
import io
import base64
from datetime import datetime, timedelta
from flask import Flask, render_template, request, jsonify, send_file
import pdfplumber
from docx import Document
import anthropic

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize Anthropic client - reads from ANTHROPIC_API_KEY env var
# You can also create a .env file or set it in your system environment
def get_anthropic_client():
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        # Try reading from config file
        config_path = os.path.join(os.path.dirname(__file__), 'config.json')
        if os.path.exists(config_path):
            with open(config_path) as f:
                config = json.load(f)
                api_key = config.get('anthropic_api_key')
    
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY not found. Set it as environment variable or in config.json")
    
    return anthropic.Anthropic(api_key=api_key)

# QuickBooks CSV fields - * indicates mandatory
QB_FIELDS = [
    '*InvoiceNo',      # Mandatory
    '*Customer',       # Mandatory
    '*InvoiceDate',    # Mandatory
    '*DueDate',        # Mandatory
    'Terms',
    'Location',
    'Memo',
    'Item(Product/Service)',
    'ItemDescription',
    'ItemQuantity',
    'ItemRate',
    '*ItemAmount',     # Mandatory
    '*ItemTaxCode',    # Mandatory
    'ItemTaxAmount'
]

EXTRACTION_PROMPT = """You are an expert at extracting invoice data. Analyze this invoice and extract ALL information accurately.

Return a JSON object with this exact structure:
{
    "invoice_no": "the invoice number exactly as shown (e.g., '25C-300 (090)', '25-086-5', etc.)",
    "customer": "the customer/client name - look for company names, 'Bill To', 'Attention' fields",
    "invoice_date": "date in DD/MM/YYYY format",
    "due_date": "due date in DD/MM/YYYY format if shown, otherwise null",
    "terms": "payment terms if shown (e.g., 'Net 30'), otherwise null",
    "subtotal": numeric subtotal before tax,
    "gst_amount": numeric GST/tax amount,
    "total": numeric total including tax,
    "line_items": [
        {
            "item": "item/service name (e.g., 'Labor', 'Materials', 'Paint', 'Skim Coat Application', 'Painting')",
            "description": "detailed description of the work/item",
            "quantity": numeric quantity,
            "rate": numeric unit rate/price,
            "amount": numeric line total
        }
    ]
}

IMPORTANT EXTRACTION RULES:
1. Invoice Number: Extract the EXACT invoice number as shown. Look for patterns like "INVOICE #", "OUR INVOICE", "Invoice No."
2. Customer: Look for the company/person being billed. Check "Bill To", "Attention:", or the main recipient name at the top. This is CRITICAL - never leave customer blank.
3. Dates: Convert all dates to DD/MM/YYYY format.
4. Line Items: Extract EVERY line item separately. Don't combine them.
5. For labor items, include hours in quantity and hourly rate.
6. GST is typically 5% in Canada.
7. All monetary values should be numbers without currency symbols.

Return ONLY the JSON object, no other text."""


def extract_text_from_pdf(filepath):
    """Extract text content from PDF"""
    text = ""
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
            
            # Also extract tables
            tables = page.extract_tables()
            for table in tables:
                if table:
                    for row in table:
                        row_text = [str(cell) if cell else "" for cell in row]
                        text += " | ".join(row_text) + "\n"
    return text


def extract_text_from_docx(filepath):
    """Extract text content from Word document"""
    doc = Document(filepath)
    text = ""
    
    for para in doc.paragraphs:
        text += para.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text += " | ".join(row_text) + "\n"
    
    return text


def extract_invoice_with_ai(text, filename):
    """Use Claude AI to extract invoice data from text"""
    try:
        client = get_anthropic_client()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[
                {
                    "role": "user",
                    "content": f"{EXTRACTION_PROMPT}\n\nINVOICE TEXT:\n{text}\n\nFILENAME: {filename}"
                }
            ]
        )
        
        response_text = message.content[0].text.strip()
        
        # Clean up response - remove markdown code blocks if present
        if response_text.startswith("```"):
            response_text = re.sub(r'^```json?\s*', '', response_text)
            response_text = re.sub(r'\s*```$', '', response_text)
        
        # Parse JSON
        invoice_data = json.loads(response_text)
        invoice_data['raw_text'] = text
        
        return invoice_data
        
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print(f"Response was: {response_text}")
        return {
            'invoice_no': None,
            'customer': None,
            'invoice_date': None,
            'due_date': None,
            'terms': None,
            'subtotal': None,
            'gst_amount': None,
            'total': None,
            'line_items': [],
            'raw_text': text,
            'error': f"Failed to parse AI response: {str(e)}"
        }
    except Exception as e:
        print(f"AI extraction error: {e}")
        return {
            'invoice_no': None,
            'customer': None,
            'invoice_date': None,
            'due_date': None,
            'terms': None,
            'subtotal': None,
            'gst_amount': None,
            'total': None,
            'line_items': [],
            'raw_text': text,
            'error': str(e)
        }


def convert_to_quickbooks_csv(invoices):
    """Convert extracted invoice data to QuickBooks CSV format"""
    rows = []
    
    for inv in invoices:
        invoice_no = inv.get('invoice_no', '') or ''
        customer = inv.get('customer', '') or ''
        invoice_date = inv.get('invoice_date', '') or ''
        
        # Calculate due date if not provided
        due_date = inv.get('due_date', '') or ''
        if not due_date and invoice_date:
            try:
                dt = datetime.strptime(invoice_date, '%d/%m/%Y')
                due_dt = dt + timedelta(days=30)
                due_date = due_dt.strftime('%d/%m/%Y')
            except:
                due_date = invoice_date
        
        terms = inv.get('terms', '') or 'Net 30'
        line_items = inv.get('line_items', []) or []
        gst_total = inv.get('gst_amount', 0) or 0
        
        if not line_items:
            # Create a single line item from total
            subtotal = inv.get('subtotal') or inv.get('total')
            
            if subtotal:
                rows.append({
                    '*InvoiceNo': invoice_no,
                    '*Customer': customer,
                    '*InvoiceDate': invoice_date,
                    '*DueDate': due_date,
                    'Terms': terms,
                    'Location': '',
                    'Memo': '',
                    'Item(Product/Service)': 'Services',
                    'ItemDescription': '',
                    'ItemQuantity': 1,
                    'ItemRate': subtotal,
                    '*ItemAmount': subtotal,
                    '*ItemTaxCode': 'GST',
                    'ItemTaxAmount': gst_total
                })
        else:
            # Calculate total amount for GST distribution
            total_amount = sum(float(item.get('amount', 0) or 0) for item in line_items)
            
            # Create rows for each line item
            is_first = True
            for item in line_items:
                item_amount = float(item.get('amount', 0) or 0)
                
                # Distribute GST proportionally across items
                if total_amount > 0 and item_amount > 0 and gst_total:
                    item_gst = round(float(gst_total) * (item_amount / total_amount), 2)
                else:
                    item_gst = ''
                
                row = {
                    '*InvoiceNo': invoice_no if is_first else '',
                    '*Customer': customer if is_first else '',
                    '*InvoiceDate': invoice_date if is_first else '',
                    '*DueDate': due_date if is_first else '',
                    'Terms': terms if is_first else '',
                    'Location': '',
                    'Memo': '',
                    'Item(Product/Service)': item.get('item', '') or '',
                    'ItemDescription': item.get('description', '') or '',
                    'ItemQuantity': item.get('quantity', '') or '',
                    'ItemRate': item.get('rate', '') or '',
                    '*ItemAmount': item_amount,
                    '*ItemTaxCode': 'GST',
                    'ItemTaxAmount': item_gst
                }
                rows.append(row)
                is_first = False
    
    return rows


@app.route('/')
def index():
    return render_template('index.html', qb_fields=QB_FIELDS)


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and AI extraction"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = file.filename
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    try:
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        elif filename.lower().endswith('.docx'):
            text = extract_text_from_docx(filepath)
        else:
            os.remove(filepath)
            return jsonify({'error': 'Unsupported file type. Please upload PDF or DOCX files.'}), 400
        
        # Use AI to extract invoice data
        invoice_data = extract_invoice_with_ai(text, filename)
        
        # Clean up
        os.remove(filepath)
        
        return jsonify({
            'success': True,
            'data': invoice_data,
            'filename': filename
        })
    
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({'error': str(e)}), 500


@app.route('/convert', methods=['POST'])
def convert_invoices():
    """Convert edited invoice data to QuickBooks CSV"""
    try:
        invoices = request.json.get('invoices', [])
        if not invoices:
            return jsonify({'error': 'No invoice data provided'}), 400
        
        csv_rows = convert_to_quickbooks_csv(invoices)
        
        # Create CSV in memory
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=QB_FIELDS)
        writer.writeheader()
        writer.writerows(csv_rows)
        
        csv_content = output.getvalue()
        
        return jsonify({
            'success': True,
            'csv': csv_content,
            'row_count': len(csv_rows)
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download', methods=['POST'])
def download_csv():
    """Generate and download the CSV file"""
    try:
        invoices = request.json.get('invoices', [])
        csv_rows = convert_to_quickbooks_csv(invoices)
        
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=QB_FIELDS)
        writer.writeheader()
        writer.writerows(csv_rows)
        
        # Convert to bytes
        mem = io.BytesIO()
        mem.write(output.getvalue().encode('utf-8'))
        mem.seek(0)
        
        return send_file(
            mem,
            mimetype='text/csv',
            as_attachment=True,
            download_name='quickbooks_invoices.csv'
        )
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5005)
